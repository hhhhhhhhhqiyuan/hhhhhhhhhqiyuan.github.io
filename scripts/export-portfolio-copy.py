"""Export rendered portfolio copy to an editable Word document."""
import json
from pathlib import Path
from docx import Document
from docx.shared import Mm, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT.parent / 'portfolio-deliverables/2026-09-06'
OUTPUT.mkdir(parents=True, exist_ok=True)
pages = json.loads((ROOT/'tmp/webp-optimization/qa/copy.json').read_text())
doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Mm(210), Mm(297)
sec.top_margin, sec.bottom_margin = Mm(19), Mm(19)
sec.left_margin, sec.right_margin = Mm(22), Mm(22)

for name in ['Normal','Title','Subtitle','Heading 1','Heading 2','Heading 3','Heading 4']:
    style = doc.styles[name]
    style.font.name = 'Arial'
    style.font.color.rgb = RGBColor(0,0,0)
    fonts=style.element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn('w:eastAsia'),'Heiti SC')
    for attr in ['asciiTheme','hAnsiTheme','eastAsiaTheme','cstheme']:
        fonts.attrib.pop(qn('w:'+attr),None)
    for border in style.element.findall('.//'+qn('w:pBdr')):
        border.getparent().remove(border)
normal=doc.styles['Normal']
normal.font.size=Pt(10.5)
normal.paragraph_format.line_spacing=1.18
normal.paragraph_format.space_after=Pt(4)
normal.paragraph_format.widow_control=True
for name,size in [('Title',24),('Heading 1',19),('Heading 2',15),('Heading 3',12),('Heading 4',11)]:
    style=doc.styles[name]
    style.font.size=Pt(size)
    style.paragraph_format.space_before=Pt(13)
    style.paragraph_format.space_after=Pt(7)
    style.paragraph_format.keep_with_next=True
    style.paragraph_format.keep_together=True

doc.add_paragraph('Nikki 作品集文案修改稿',style='Title')
doc.add_paragraph('首页 关于我及八个项目',style='Subtitle')
doc.add_paragraph('这里按当前网站顺序整理了标题、正文、数据标签和图注。你可以直接修改原文，也可以用 Word 的修订或批注功能留下想法；修改完成后，将这份文档发回即可对应更新网站。')
doc.add_paragraph('此稿保留网页原文，尚未润色。设计稿、品牌展示图和嵌入式网页内部的文字不纳入这次文案修改；正文中显示的项目数据也保留原值。')
doc.add_paragraph('导出日期 2026 年 9 月 6 日')

all_text=['Nikki 作品集文案修改稿','导出日期 2026 年 9 月 6 日','']
count=0
for index,page in enumerate(pages,1):
    title=doc.add_paragraph(f'{index:02d} {page["label"]}',style='Heading 1')
    if index>1:title.paragraph_format.space_before=Pt(24)
    all_text.extend(['',f'{index:02d} {page["label"]}',''])
    for b in page['blocks']:
        text=b['text']
        # Preserve wording; restore the visual line breaks flattened during capture.
        # Slash-separated interface labels remain unchanged for accurate editing.
        kind=b['type']
        if kind in ['h1','h2']: style='Heading 2'
        elif kind in ['h3','h4','h5','h6']: style='Heading 3'
        else: style='Normal'
        para=doc.add_paragraph(text,style=style)
        if kind=='dt':
            for run in para.runs:run.bold=True
        if kind=='figcaption':
            for run in para.runs:run.italic=False
        count+=1
        all_text.append(text)

footer=sec.footer.paragraphs[0]
footer.alignment=2
run=footer.add_run('Nikki 作品集文案  ')
run.font.size=Pt(9)
field=OxmlElement('w:fldSimple');field.set(qn('w:instr'),'PAGE');footer._p.append(field)
doc.core_properties.title='Nikki 作品集文案修改稿'
doc.core_properties.author='Nikki'
doc.core_properties.subject='现有网页文案 可编辑修改稿'
for paragraph in list(doc.paragraphs)+list(sec.footer.paragraphs):
    for run in paragraph.runs:
        fonts=run._element.get_or_add_rPr().get_or_add_rFonts()
        fonts.set(qn('w:eastAsia'),'Heiti SC')
dest=OUTPUT/'Nikki-作品集文案-可编辑.docx'
doc.save(dest)
(OUTPUT/'Nikki-作品集文案-纯文字.txt').write_text('\n'.join(all_text),encoding='utf-8')
print(json.dumps({'docx':str(dest),'projects':len(pages),'text_blocks':count},ensure_ascii=False))
